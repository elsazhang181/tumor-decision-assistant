#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import glob
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 查找 Markdown 文件
md_files = glob.glob('/workspace/projects/public/*.md')
print(f'找到 Markdown 文件：{md_files}')

if not md_files:
    print('未找到 Markdown 文件！')
    exit(1)

md_path = md_files[0]
print(f'使用文件：{md_path}')

# 读取 Markdown 文件
with open(md_path, 'r', encoding='utf-8') as f:
    md_content = f.read()

# 创建 Word 文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

lines = md_content.split('\n')

for line in lines:
    # 标题
    if line.startswith('# ') and not line.startswith('## '):
        heading = doc.add_heading(line[2:], level=1)
        for run in heading.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    elif line.startswith('## ') and not line.startswith('### '):
        heading = doc.add_heading(line[3:], level=2)
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    elif line.startswith('### ') and not line.startswith('#### '):
        heading = doc.add_heading(line[4:], level=3)
        for run in heading.runs:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x3d, 0x3d, 0x3d)
    elif line.startswith('#### '):
        heading = doc.add_heading(line[5:], level=4)
        for run in heading.runs:
            run.font.size = Pt(11)
    # 表格行
    elif line.startswith('|') and line.endswith('|') and not line.startswith('|---'):
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            p = doc.add_paragraph()
            p.add_run(' | '.join(cells)).font.size = Pt(10)
    # 分隔线
    elif line.startswith('---'):
        p = doc.add_paragraph()
        p.add_run('─' * 50).font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    # 列表项
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)
    # 代码块标记
    elif line.startswith('```'):
        continue
    # 普通文本
    elif line.strip():
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)

# 保存 Word 文件
output_path = '/workspace/projects/public/健康就医决策助手 - 设计理念与部署文档.docx'
doc.save(output_path)
print(f'Word 文档已生成：{output_path}')
